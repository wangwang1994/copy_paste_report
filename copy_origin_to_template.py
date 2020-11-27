import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
filename=input('请输入原始word报告，必须使用word另存为docx格式：')
origin_excel=input('请输入原始excel文件名:')
baogaobianhao=input('请输入报告编号：')
yangpinbianhao=input('请输入样品编号：')
gongsimingcheng=input('请输入抽检公司名称：')
doc=Document(filename)
print(os.getcwd())
# for table in doc.tables:
#     print(table)
# print(doc.tables[2].cell(0,2).text)
# def copy_table_after(table, paragraph):
#     tbl, p = table._tbl, paragraph._p
#     new_tbl = deepcopy(tbl)
#     p.addnext(new_tbl)



doc.tables[2].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[2].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[4].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[4].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[6].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[6].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[8].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[8].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.paragraphs[0].text='报告编号：'+baogaobianhao
paragraph=doc.paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.tables[3].cell(10,1).text="检验结果见第3页"
paragraph=doc.tables[3].cell(10,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)

doc.tables[3].cell(6,1).text='郑州市生态环境局'
paragraph=doc.tables[3].cell(6,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[3].cell(5,3).text='朱永宏'
paragraph=doc.tables[3].cell(5,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[3].cell(6,3).text='--'
paragraph=doc.tables[3].cell(6,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

count=0
text=[]
for par in doc.paragraphs:
    count+=1
    # print(par.text)
    # print(count)
    text.append(par.text)
# print(doc.paragraphs[60].text)
# print(doc.paragraphs[61].text)
# print(text)
for item in text:
    if item.startswith('检验日期')==True:
        jianyanshijian=item
    if item.startswith('检验地点')==True:
        jianyandidian=item
print(jianyanshijian)
print(jianyandidian)

jianyanshijian_1=jianyanshijian.replace('检验日期：','')
doc.tables[3].cell(4,3).text=jianyanshijian_1
paragraph=doc.tables[3].cell(4,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

print(doc.tables[7].cell(4,1).text)
doc.tables[7].cell(4,1).text="--"
paragraph=doc.tables[7].cell(4,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(4,7).text='--'
paragraph=doc.tables[7].cell(4,6).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(6,7).text='--'
paragraph=doc.tables[7].cell(6,7).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

calid1=doc.tables[7].cell(14,5).text
cvn1=doc.tables[7].cell(14,8).text
calid2=doc.tables[7].cell(16,5).text
cvn2=doc.tables[7].cell(16,8).text

# copy_table_after(doc.tables[0],doc.paragraphs[1])
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
table=doc.tables[7]
row1 = doc.tables[7].rows[14]
row2 = doc.tables[7].rows[15]
row3= doc.tables[7].rows[16]
row4 = doc.tables[7].rows[17]

remove_row(table, row1)
remove_row(table, row2)
remove_row(table, row3)
remove_row(table, row4)
count=0
# for par in doc.paragraphs:
#     count+=1
#     print(par.text)
#     print(count)
# print(doc.paragraphs[60].text)
# print(doc.paragraphs[61].text)

doc.save(baogaobianhao+'.docx')

import os,sys
import openpyxl
from openpyxl.drawing.image import Image
from zipfile import ZipFile

wb_origin = openpyxl.load_workbook(origin_excel)  # Add file name
sheet_canshu = wb_origin["参数"]  # Add Sheet name
# 由于合并的单元格无法在以下的循环中进行使用，因此需要在这里将需要循环部分的
# 单元格进行拆分，在最后保存之前进行合并就可以了
for i in range(4,16):
    sheet_canshu.unmerge_cells(start_row=i, start_column=2, end_row=i, end_column=3)
    sheet_canshu.unmerge_cells(start_row=i, start_column=4, end_row=i, end_column=5)
    sheet_canshu.unmerge_cells(start_row=i, start_column=6, end_row=i, end_column=7)
    sheet_canshu.unmerge_cells(start_row=i, start_column=8, end_row=i, end_column=9)
sheet_yuanshijilu=wb_origin['原始记录']
sheet_yuanshijilu.unmerge_cells('C5:D5')

# File to be pasted into
template = openpyxl.load_workbook("轻型汽油车原始记录模板.xlsx")  # Add file name
temp_sheet_canshu = template["参数"]  # Add Sheet name
temp_sheet_yuanshijilu=template['原始记录']



def copyRange(startCol,startRow,endCol,endRow,sheet):
    rangeSelected = []
    # Loops through selected Rows
    for i in range(startRow, endRow + 1, 1):
        # Appends the row to a RowSelected list
        rowSelected = []
        for j in range(startCol, endCol + 1, 1):
            rowSelected.append(sheet.cell(row=i, column=j).value)
        # Adds the RowSelected List and nests inside the rangeSelected
        rangeSelected.append(rowSelected)
    return rangeSelected


# Paste range
# Paste data from copyRange into template sheet
def pasteRange(startCol, startRow, endCol, endRow, sheetReceiving, copiedData):
    countRow = 0
    for i in range(startRow, endRow + 1, 1):
        countCol = 0
        for j in range(startCol, endCol + 1, 1):
            sheetReceiving.cell(row=i, column=j).value = copiedData[countRow][countCol]
            countCol += 1
        countRow += 1

def createData():
    print("Processing...")
    selectedRange_canshu = copyRange(2, 4, 9, 15, sheet_canshu)
    pastingRange_canshu = pasteRange(2, 4, 9, 15, temp_sheet_canshu,selectedRange_canshu)
    selectedRange_yuanshijilu1=copyRange(3,5,4,5,sheet_yuanshijilu)
    pastingRange_yuanshijilu1=pasteRange(3,5,4,5,temp_sheet_yuanshijilu,selectedRange_yuanshijilu1)
    selectedRange_yuanshijilu2=copyRange(4,9,4,9,sheet_yuanshijilu)
    pastingRange_yuanshijilu2=pasteRange(4,10,4,10,temp_sheet_yuanshijilu,selectedRange_yuanshijilu2)
    selectedRange_yuanshijilu3 = copyRange(4, 11, 4, 11, sheet_yuanshijilu)
    pastingRange_yuanshijilu3 = pasteRange(4, 12, 4, 12, temp_sheet_yuanshijilu, selectedRange_yuanshijilu3)

    for i in range(4, 16):
        temp_sheet_canshu.merge_cells(start_row=i, start_column=2, end_row=i, end_column=3)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=4, end_row=i, end_column=5)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=6, end_row=i, end_column=7)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=8, end_row=i, end_column=9)
    temp_sheet_yuanshijilu.merge_cells('C5:D5')
    temp_sheet_yuanshijilu['F43']='报告编号：'+baogaobianhao
    temp_sheet_yuanshijilu['B94']='外观检验照片见'+baogaobianhao+'#光盘 文件夹'
    temp_sheet_yuanshijilu['C9']=yangpinbianhao
    temp_sheet_yuanshijilu['A25']='4.'+jianyanshijian
    temp_sheet_yuanshijilu['A26']='5.'+jianyandidian+gongsimingcheng
    temp_sheet_yuanshijilu['G102']=calid1
    temp_sheet_yuanshijilu['I102']=cvn1
    temp_sheet_yuanshijilu['G104']=calid2
    temp_sheet_yuanshijilu['I104']=cvn2
    template.save("参数页复制后.xlsx")
    print("Range copied and pasted!")
createData()
wb_xiugaihou=openpyxl.load_workbook('参数页复制后.xlsx')
ws_suicheqingdan=wb_xiugaihou['随车清单']

zip = ZipFile(origin_excel)
zip.extractall()
try:
    os.rename('xl/media/image1.jpeg','xl/media/image1.png')
except:
    pass
img1 = Image('xl/media/image1.png')
img1.height=500
img1.width=500
img1.anchor='A3'
ws_suicheqingdan.add_image(img1)
try:
    os.rename('xl/media/image2.jpeg', 'xl/media/image2.png')
except:
    print('没有第二张清单')
try:
    img2=Image('xl/media/image2.png')
    img2.height = 500
    img2.width = 500
    img2.anchor = 'A18'
    ws_suicheqingdan.add_image(img2)
except FileNotFoundError:
    print('没有第二张清单')
wb_xiugaihou.save(baogaobianhao+'.xlsx')
os.remove('xl/media/image1.png')
try:
    os.remove('xl/media/image2.png')
except:
    pass


