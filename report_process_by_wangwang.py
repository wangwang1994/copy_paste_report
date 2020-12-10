import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
def get_info(file_name):
    os.chdir(file_name)
    file_name=file_name
    baogaobianhao=open('报告编号.txt')
    baogaobianhao_list=[]
    for line in baogaobianhao.readlines():
        line=line.rstrip('\n')
        baogaobianhao_list.append(line)
    yangpinbianhao = open('样品编号.txt')
    yangpinbianhao_list = []
    for line in yangpinbianhao.readlines():
        line = line.rstrip('\n')
        yangpinbianhao_list.append(line)

    first_file_list=os.listdir()
    first_file_list.remove('报告编号.txt')
    first_file_list.remove('样品编号.txt')
    try:
        first_file_list.remove('.DS_Store')
    except:
        pass
    docx_dict={}
    xlsx_dict={}
    for i in range(len(first_file_list)):
        second_file_name=file_name+'/'+first_file_list[i]
        os.chdir(second_file_name)
        second_file_list=os.listdir()
        for f in second_file_list:
            if f.endswith('.docx'):
                docx_dict[first_file_list[i]]=f
            if f.endswith(('.xlsx')):
                xlsx_dict[first_file_list[i]]=f
    # print(docx_dict)
    # print(xlsx_dict)
        # print(second_file_list)
        # print(second_file_name)
    baogaobianhao_dict={}
    yangpinbianhao_dict={}
    for i in range(len(first_file_list)):
        baogaobianhao_dict[first_file_list[i]]=baogaobianhao_list[i]
    for j in range(len(first_file_list)):
        yangpinbianhao_dict[first_file_list[j]]=yangpinbianhao_list[j]



    return file_name,first_file_list,baogaobianhao_dict,yangpinbianhao_dict,docx_dict,xlsx_dict

file_name,chexing_list,baogaobianhao_dict,yangpinbianhao_dict,docx_dict,xlsx_dict=get_info('/Users/wangwang/Desktop/测试文件夹')
print(file_name)
print(chexing_list)
print(baogaobianhao_dict)
print(yangpinbianhao_dict)
print(docx_dict)
print(xlsx_dict)




# filename = input('请输入word报告名，以docx结尾：')
filename = docx_dict[chexing_list[0]]
print(filename)
# origin_excel = input('请输入原始excel的名称：')
origin_excel = xlsx_dict[chexing_list[0]]
print(origin_excel)
# baogaobianhao = input('请输入报告编号：')
baogaobianhao = baogaobianhao_dict[chexing_list[0]]
print(baogaobianhao)
# yangpinbianhao = input('请输入样品编号：')
yangpinbianhao = yangpinbianhao_dict[chexing_list[0]]
print(yangpinbianhao)
gongsimingcheng = input('请输入抽检公司名称：')

def copy_paste_docx(doc_file_name, baogaobianhao_name):
    doc = Document(doc_file_name)
    doc_canshu = Document('参数确认表_模版.docx')
    doc.tables[2].cell(0, 2).text = '报告编号：' + baogaobianhao_name
    paragraph = doc.tables[2].cell(0, 2).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.tables[4].cell(0, 2).text = '报告编号：' + baogaobianhao
    paragraph = doc.tables[4].cell(0, 2).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.tables[6].cell(0, 2).text = '报告编号：' + baogaobianhao
    paragraph = doc.tables[6].cell(0, 2).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.tables[8].cell(0, 2).text = '报告编号：' + baogaobianhao
    paragraph = doc.tables[8].cell(0, 2).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.paragraphs[0].text = '报告编号：' + baogaobianhao
    paragraph = doc.paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.tables[3].cell(10, 1).text = "检验结果见第3页"
    paragraph = doc.tables[3].cell(10, 1).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)

    doc.tables[3].cell(6, 1).text = '郑州市生态环境局'
    paragraph = doc.tables[3].cell(6, 1).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.tables[3].cell(5, 3).text = '朱永宏'
    paragraph = doc.tables[3].cell(5, 3).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.tables[3].cell(6, 3).text = '--'
    paragraph = doc.tables[3].cell(6, 3).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    count = 0
    text = []
    for par in doc.paragraphs:
        count += 1
        # print(par.text)
        # print(count)
        text.append(par.text)
    # print(doc.paragraphs[60].text)
    # print(doc.paragraphs[61].text)
    # print(text)
    for item in text:
        if item.startswith('检验日期') == True:
            global jianyanshijian
            jianyanshijian = item

        if item.startswith('检验地点') == True:
            global jianyandidian
            jianyandidian = item

    print(jianyanshijian)
    print(jianyandidian)

    jianyanshijian_1 = jianyanshijian.replace('检验日期：', '')
    doc.tables[3].cell(4, 3).text = jianyanshijian_1
    paragraph = doc.tables[3].cell(4, 3).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(doc.tables[7].cell(4, 1).text)
    doc.tables[7].cell(4, 1).text = "--"
    paragraph = doc.tables[7].cell(4, 1).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.tables[7].cell(4, 7).text = '--'
    paragraph = doc.tables[7].cell(4, 6).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.tables[7].cell(6, 7).text = '--'
    paragraph = doc.tables[7].cell(6, 7).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    global calid1
    global calid2
    global cvn1
    global cvn2
    calid1 = doc.tables[7].cell(14, 5).text
    cvn1 = doc.tables[7].cell(14, 8).text
    calid2 = doc.tables[7].cell(16, 5).text
    cvn2 = doc.tables[7].cell(16, 8).text

    # copy_table_after(doc.tables[0],doc.paragraphs[1])
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    table = doc.tables[7]
    row1 = doc.tables[7].rows[14]
    row2 = doc.tables[7].rows[15]
    row3 = doc.tables[7].rows[16]
    row4 = doc.tables[7].rows[17]

    remove_row(table, row1)
    remove_row(table, row2)
    remove_row(table, row3)
    remove_row(table, row4)
    count = 0
    # for par in doc.paragraphs:
    #     count+=1
    #     print(par.text)
    #     print(count)
    # print(doc.paragraphs[60].text)
    # print(doc.paragraphs[61].text)
    doc.save(baogaobianhao_name + '.docx')
    canshuqueren_table = doc.tables[7]

    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def get_paragraph(paras, text):
        for para in paras:
            if text in para.text:
                return para
        raise KeyError("The text cannot be found anywhere in the document")

    para1 = get_paragraph(doc_canshu.paragraphs, '参数确认表')
    move_table_after(canshuqueren_table, para1)

    canshuqueren_row1 = doc_canshu.tables[0].rows[0]
    canshuqueren_row2 = doc_canshu.tables[0].rows[10]
    canshuqueren_row3 = doc_canshu.tables[0].rows[11]
    canshuqueren_row4 = doc_canshu.tables[0].rows[12]
    canshuqueren_row5 = doc_canshu.tables[0].rows[13]

    remove_row(canshuqueren_table, canshuqueren_row1)
    remove_row(canshuqueren_table, canshuqueren_row2)
    remove_row(canshuqueren_table, canshuqueren_row3)
    remove_row(canshuqueren_table, canshuqueren_row4)
    remove_row(canshuqueren_table, canshuqueren_row5)

    doc_canshu.save(''+baogaobianhao+'参数确认表.docx')
copy_paste_docx(filename,baogaobianhao)

print('----------------------以上完成了DOCX函数的封装---------------------')



from copy import deepcopy
import os

filename=input('请输入word报告名，以docx结尾：')
origin_excel=input('请输入原始excel的名称：')
baogaobianhao=input('请输入报告编号：')
yangpinbianhao=input('请输入样品编号：')
gongsimingcheng=input('请输入抽检公司名称：')
doc=Document(filename)


doc_canshu=Document('参数确认表_模版.docx')
doc.tables[2].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[2].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[4].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[4].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[6].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[6].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[8].cell(0,2).text='报告编号：'+baogaobianhao
paragraph=doc.tables[8].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.paragraphs[0].text='报告编号：'+baogaobianhao
paragraph=doc.paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.tables[3].cell(10,1).text="检验结果见第3页"
paragraph=doc.tables[3].cell(10,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)

doc.tables[3].cell(6,1).text='郑州市生态环境局'
paragraph=doc.tables[3].cell(6,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[3].cell(5,3).text='朱永宏'
paragraph=doc.tables[3].cell(5,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[3].cell(6,3).text='--'
paragraph=doc.tables[3].cell(6,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

count=0
text=[]
for par in doc.paragraphs:
    count+=1
    # print(par.text)
    # print(count)
    text.append(par.text)
# print(doc.paragraphs[60].text)
# print(doc.paragraphs[61].text)
# print(text)
for item in text:
    if item.startswith('检验日期')==True:
        jianyanshijian=item
    if item.startswith('检验地点')==True:
        jianyandidian=item
print(jianyanshijian)
print(jianyandidian)

jianyanshijian_1=jianyanshijian.replace('检验日期：','')
doc.tables[3].cell(4,3).text=jianyanshijian_1
paragraph=doc.tables[3].cell(4,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

print(doc.tables[7].cell(4,1).text)
doc.tables[7].cell(4,1).text="--"
paragraph=doc.tables[7].cell(4,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(4,7).text='--'
paragraph=doc.tables[7].cell(4,6).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(6,7).text='--'
paragraph=doc.tables[7].cell(6,7).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

calid1=doc.tables[7].cell(14,5).text
cvn1=doc.tables[7].cell(14,8).text
calid2=doc.tables[7].cell(16,5).text
cvn2=doc.tables[7].cell(16,8).text

# copy_table_after(doc.tables[0],doc.paragraphs[1])
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
table=doc.tables[7]
row1 = doc.tables[7].rows[14]
row2 = doc.tables[7].rows[15]
row3= doc.tables[7].rows[16]
row4 = doc.tables[7].rows[17]

remove_row(table, row1)
remove_row(table, row2)
remove_row(table, row3)
remove_row(table, row4)
count=0
# for par in doc.paragraphs:
#     count+=1
#     print(par.text)
#     print(count)
# print(doc.paragraphs[60].text)
# print(doc.paragraphs[61].text)
doc.save(baogaobianhao+'.docx')
canshuqueren_table=doc.tables[7]


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
def get_paragraph(paras, text):
    for para in paras:
        if text in para.text:
            return para
    raise KeyError("The text cannot be found anywhere in the document")
para1=get_paragraph(doc_canshu.paragraphs,'参数确认表')
move_table_after(canshuqueren_table,para1)


canshuqueren_row1 = doc_canshu.tables[0].rows[0]
canshuqueren_row2 = doc_canshu.tables[0].rows[10]
canshuqueren_row3 = doc_canshu.tables[0].rows[11]
canshuqueren_row4 = doc_canshu.tables[0].rows[12]
canshuqueren_row5 = doc_canshu.tables[0].rows[13]

remove_row(canshuqueren_table, canshuqueren_row1)
remove_row(canshuqueren_table, canshuqueren_row2)
remove_row(canshuqueren_table, canshuqueren_row3)
remove_row(canshuqueren_table, canshuqueren_row4)
remove_row(canshuqueren_table, canshuqueren_row5)

doc_canshu.save('参数确认表.docx')




import os
import openpyxl
from openpyxl.drawing.image import Image
from zipfile import ZipFile

wb_origin = openpyxl.load_workbook(origin_excel)  # Add file name
sheet_canshu = wb_origin["参数"]  # Add Sheet name
# 由于合并的单元格无法在以下的循环中进行使用，因此需要在这里将需要循环部分的
# 单元格进行拆分，在最后保存之前进行合并就可以了
for i in range(4,16):
    sheet_canshu.unmerge_cells(start_row=i, start_column=2, end_row=i, end_column=3)
    sheet_canshu.unmerge_cells(start_row=i, start_column=4, end_row=i, end_column=5)
    sheet_canshu.unmerge_cells(start_row=i, start_column=6, end_row=i, end_column=7)
    sheet_canshu.unmerge_cells(start_row=i, start_column=8, end_row=i, end_column=9)
sheet_yuanshijilu=wb_origin['原始记录']
sheet_yuanshijilu.unmerge_cells('C5:D5')

# File to be pasted into
template = openpyxl.load_workbook("轻型汽油车原始记录模板.xlsx")  # Add file name这里是黏贴的模版
temp_sheet_canshu = template["参数"]  # Add Sheet name
temp_sheet_yuanshijilu=template['原始记录']



def copyRange(startCol,startRow,endCol,endRow,sheet):
    rangeSelected = []
    # Loops through selected Rows
    for i in range(startRow, endRow + 1, 1):
        # Appends the row to a RowSelected list
        rowSelected = []
        for j in range(startCol, endCol + 1, 1):
            rowSelected.append(sheet.cell(row=i, column=j).value)
        # Adds the RowSelected List and nests inside the rangeSelected
        rangeSelected.append(rowSelected)
    return rangeSelected


# Paste range
# Paste data from copyRange into template sheet
def pasteRange(startCol, startRow, endCol, endRow, sheetReceiving, copiedData):
    countRow = 0
    for i in range(startRow, endRow + 1, 1):
        countCol = 0
        for j in range(startCol, endCol + 1, 1):
            sheetReceiving.cell(row=i, column=j).value = copiedData[countRow][countCol]
            countCol += 1
        countRow += 1

def createData():
    print("Processing...")
    selectedRange_canshu = copyRange(2, 4, 9, 15, sheet_canshu)
    pastingRange_canshu = pasteRange(2, 4, 9, 15, temp_sheet_canshu,selectedRange_canshu)
    selectedRange_yuanshijilu1=copyRange(3,5,4,5,sheet_yuanshijilu)
    pastingRange_yuanshijilu1=pasteRange(3,5,4,5,temp_sheet_yuanshijilu,selectedRange_yuanshijilu1)
    selectedRange_yuanshijilu2=copyRange(4,9,4,9,sheet_yuanshijilu)
    pastingRange_yuanshijilu2=pasteRange(4,10,4,10,temp_sheet_yuanshijilu,selectedRange_yuanshijilu2)
    selectedRange_yuanshijilu3 = copyRange(4, 11, 4, 11, sheet_yuanshijilu)
    pastingRange_yuanshijilu3 = pasteRange(4, 12, 4, 12, temp_sheet_yuanshijilu, selectedRange_yuanshijilu3)

    for i in range(4, 16):
        temp_sheet_canshu.merge_cells(start_row=i, start_column=2, end_row=i, end_column=3)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=4, end_row=i, end_column=5)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=6, end_row=i, end_column=7)
        temp_sheet_canshu.merge_cells(start_row=i, start_column=8, end_row=i, end_column=9)
    temp_sheet_yuanshijilu.merge_cells('C5:D5')
    temp_sheet_yuanshijilu['F43']='报告编号：'+baogaobianhao
    temp_sheet_yuanshijilu['B94']='外观检验照片见'+baogaobianhao+'#光盘 文件夹'
    temp_sheet_yuanshijilu['C9']=yangpinbianhao
    temp_sheet_yuanshijilu['A25']='4.'+jianyanshijian
    temp_sheet_yuanshijilu['A26']='5.'+jianyandidian+gongsimingcheng
    temp_sheet_yuanshijilu['G102']=calid1
    temp_sheet_yuanshijilu['I102']=cvn1
    temp_sheet_yuanshijilu['G104']=calid2
    temp_sheet_yuanshijilu['I104']=cvn2
    template.save("参数页复制后.xlsx")
    print("Range copied and pasted!")
createData()
wb_xiugaihou=openpyxl.load_workbook('参数页复制后.xlsx')
ws_suicheqingdan=wb_xiugaihou['随车清单']

zip = ZipFile(origin_excel)
zip.extractall()
try:
    os.rename('xl/media/image1.jpeg','xl/media/image1.png')
except:
    pass
try:
    os.rename('xl/media/image1.jpg','xl/media/image1.png')
except:
    pass
img1 = Image('xl/media/image1.png')
img1.height=500
img1.width=500
img1.anchor='A3'
ws_suicheqingdan.add_image(img1)
try:
    os.rename('xl/media/image2.jpeg', 'xl/media/image2.png')
except:
    print('没有第二张清单')
try:
    img2=Image('xl/media/image2.png')
    img2.height = 500
    img2.width = 500
    img2.anchor = 'A18'
    ws_suicheqingdan.add_image(img2)
except FileNotFoundError:
    print('没有第二张清单')
wb_xiugaihou.save(baogaobianhao+'.xlsx')
os.remove('xl/media/image1.png')
try:
    os.remove('xl/media/image2.png')
except:
    pass





