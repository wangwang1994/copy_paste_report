import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
def get_info(file_name):
    os.chdir(file_name)
    file_name=file_name
    baogaobianhao=open('报告编号.txt')
    baogaobianhao_list=[]
    for line in baogaobianhao.readlines():
        line=line.rstrip('\n')
        baogaobianhao_list.append(line)
    yangpinbianhao = open('样品编号.txt')
    yangpinbianhao_list = []
    for line in yangpinbianhao.readlines():
        line = line.rstrip('\n')
        yangpinbianhao_list.append(line)

    first_file_list=os.listdir()
    first_file_list.remove('报告编号.txt')
    first_file_list.remove('样品编号.txt')
    first_file_list.remove('参数确认表_模版.docx')
    try:
        first_file_list.remove('.DS_Store')
    except:
        pass
    docx_dict={}
    xlsx_dict={}
    for i in range(len(first_file_list)):
        second_file_name=file_name+'/'+first_file_list[i]
        os.chdir(second_file_name)
        second_file_list=os.listdir()
        for f in second_file_list:
            if f.endswith('.docx'):
                docx_dict[first_file_list[i]]=f
            if f.endswith(('.xlsx')):
                xlsx_dict[first_file_list[i]]=f
    # print(docx_dict)
    # print(xlsx_dict)
        # print(second_file_list)
        # print(second_file_name)
    baogaobianhao_dict={}
    yangpinbianhao_dict={}
    for i in range(len(first_file_list)):
        baogaobianhao_dict[first_file_list[i]]=baogaobianhao_list[i]
    for j in range(len(first_file_list)):
        yangpinbianhao_dict[first_file_list[j]]=yangpinbianhao_list[j]

    os.chdir(file_name)

    return file_name,first_file_list,baogaobianhao_dict,yangpinbianhao_dict,docx_dict,xlsx_dict

file_name,chexing_list,baogaobianhao_dict,yangpinbianhao_dict,docx_dict,xlsx_dict=get_info('/Users/wangwang/Desktop/测试文件夹')
print(file_name)
print(chexing_list)
print(baogaobianhao_dict)
print(yangpinbianhao_dict)
print(docx_dict)
print(xlsx_dict)


print('完成getinfo函数后python的运行位置：')
print(os.getcwd())
gongsimingcheng = input('请输入抽检公司名称：')
# print('将python的运行位置更改到下面：')
for i in range(len(chexing_list)):
    os.chdir(file_name+'/'+chexing_list[i])
    print(os.getcwd())
    filename = docx_dict[chexing_list[i]]
    print(os.getcwd())
    print(filename)
    origin_excel = xlsx_dict[chexing_list[i]]
    print(origin_excel)
    baogaobianhao = baogaobianhao_dict[chexing_list[i]]
    print(baogaobianhao)
    yangpinbianhao = yangpinbianhao_dict[chexing_list[i]]
    print(yangpinbianhao)

    def copy_paste_docx(doc_file_name, baogaobianhao_name):
        doc = Document(doc_file_name)
        doc_canshu = Document('/Users/wangwang/Desktop/测试文件夹/参数确认表_模版.docx')
        doc.tables[2].cell(0, 2).text = '报告编号：' + baogaobianhao_name
        paragraph = doc.tables[2].cell(0, 2).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[4].cell(0, 2).text = '报告编号：' + baogaobianhao
        paragraph = doc.tables[4].cell(0, 2).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[6].cell(0, 2).text = '报告编号：' + baogaobianhao
        paragraph = doc.tables[6].cell(0, 2).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[8].cell(0, 2).text = '报告编号：' + baogaobianhao
        paragraph = doc.tables[8].cell(0, 2).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.paragraphs[0].text = '报告编号：' + baogaobianhao
        paragraph = doc.paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.tables[3].cell(10, 1).text = "检验结果见第3页"
        paragraph = doc.tables[3].cell(10, 1).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)

        doc.tables[3].cell(6, 1).text = '郑州市生态环境局'
        paragraph = doc.tables[3].cell(6, 1).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[3].cell(5, 3).text = '朱永宏'
        paragraph = doc.tables[3].cell(5, 3).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[3].cell(6, 3).text = '--'
        paragraph = doc.tables[3].cell(6, 3).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        count = 0
        text = []
        for par in doc.paragraphs:
            count += 1
            # print(par.text)
            # print(count)
            text.append(par.text)
        # print(doc.paragraphs[60].text)
        # print(doc.paragraphs[61].text)
        # print(text)
        for item in text:
            if item.startswith('检验日期') == True:
                global jianyanshijian
                jianyanshijian = item

            if item.startswith('检验地点') == True:
                global jianyandidian
                jianyandidian = item

        print(jianyanshijian)
        print(jianyandidian)

        jianyanshijian_1 = jianyanshijian.replace('检验日期：', '')
        doc.tables[3].cell(4, 3).text = jianyanshijian_1
        paragraph = doc.tables[3].cell(4, 3).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        print(doc.tables[7].cell(4, 1).text)
        doc.tables[7].cell(4, 1).text = "--"
        paragraph = doc.tables[7].cell(4, 1).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[7].cell(4, 7).text = '--'
        paragraph = doc.tables[7].cell(4, 6).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[7].cell(6, 7).text = '--'
        paragraph = doc.tables[7].cell(6, 7).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        global calid1
        global calid2
        global cvn1
        global cvn2
        calid1 = doc.tables[7].cell(14, 5).text
        cvn1 = doc.tables[7].cell(14, 8).text
        calid2 = doc.tables[7].cell(16, 5).text
        cvn2 = doc.tables[7].cell(16, 8).text
        print(calid1)
        print(calid2)
        print(cvn1)
        print(cvn2)
        # copy_table_after(doc.tables[0],doc.paragraphs[1])
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        table = doc.tables[7]
        row1 = doc.tables[7].rows[14]
        row2 = doc.tables[7].rows[15]
        row3 = doc.tables[7].rows[16]
        row4 = doc.tables[7].rows[17]

        remove_row(table, row1)
        remove_row(table, row2)
        remove_row(table, row3)
        remove_row(table, row4)
        count = 0
        # for par in doc.paragraphs:
        #     count+=1
        #     print(par.text)
        #     print(count)
        # print(doc.paragraphs[60].text)
        # print(doc.paragraphs[61].text)
        doc.save(baogaobianhao_name + '.docx')
        canshuqueren_table = doc.tables[7]

        def move_table_after(table, paragraph):
            tbl, p = table._tbl, paragraph._p
            p.addnext(tbl)

        def get_paragraph(paras, text):
            for para in paras:
                if text in para.text:
                    return para
            raise KeyError("The text cannot be found anywhere in the document")

        para1 = get_paragraph(doc_canshu.paragraphs, '参数确认表')
        move_table_after(canshuqueren_table, para1)

        canshuqueren_row1 = doc_canshu.tables[0].rows[0]
        canshuqueren_row2 = doc_canshu.tables[0].rows[10]
        canshuqueren_row3 = doc_canshu.tables[0].rows[11]
        canshuqueren_row4 = doc_canshu.tables[0].rows[12]
        canshuqueren_row5 = doc_canshu.tables[0].rows[13]

        remove_row(canshuqueren_table, canshuqueren_row1)
        remove_row(canshuqueren_table, canshuqueren_row2)
        remove_row(canshuqueren_table, canshuqueren_row3)
        remove_row(canshuqueren_table, canshuqueren_row4)
        remove_row(canshuqueren_table, canshuqueren_row5)

        doc_canshu.save(''+baogaobianhao+'参数确认表.docx')
        os.chdir(file_name)
    copy_paste_docx(filename,baogaobianhao)





print('----------------------以上完成了DOCX函数的封装---------------------')
print('完成docx处理函数后python的运行位置')
print(os.getcwd())



