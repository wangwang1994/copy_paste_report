import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = docx.Document()
paragraph=doc.add_paragraph('这是一行内容')
prior_paragraph=paragraph.insert_paragraph_before('在前面添加了一行内容，')

paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('The REAL meaning of the universe')
doc.add_heading('The role of dolphins', level=2)
doc.add_page_break()
paragraph=doc.add_paragraph('这是新一页的内容')
table=doc.add_table(rows=2,cols=2)
cell=table.cell(0,0)
cell.text='这是左上角的单元格的内容'
doc.save("test.docx")