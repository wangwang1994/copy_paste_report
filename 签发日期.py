from docx import Document
from docx.shared import Pt
doc=Document('报告轻型1.docx')
doc.tables[3].cell(11,1).text='                                      签发日期：2020年11月16日 '
paragraph=doc.tables[3].cell(11,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
doc.save('报告添加签发日期后.docx')