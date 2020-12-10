from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
name_list=[]
shijian_list=[]
f1=open('编号.txt')
f2=open('时间.txt')
for line in f1:
    line=line.rstrip('\n')
    name_list.append(line)
for line in f2:
    line = line.rstrip('\n')
    shijian_list.append(line)
# print(name_list)
# print(shijian_list)
for i in range(len(name_list)):
    doc=Document(name_list[i]+'.docx')
    doc.tables[3].cell(7,3).text=shijian_list[i]
    paragraph=doc.tables[3].cell(7,3).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.tables[3].cell(11, 1).text = '                                      签发日期：2020年11月16日 '
    paragraph = doc.tables[3].cell(11, 1).paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(10)

    doc.save(name_list[i]+'.docx')
