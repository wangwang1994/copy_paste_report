from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
# filename=
doc=Document('报告.docx')
doc_canshu=Document('参数确认表_模版.docx')
# for table in doc.tables:
#     print(table)
# print(doc.tables[2].cell(0,2).text)
# def copy_table_after(table, paragraph):
#     tbl, p = table._tbl, paragraph._p
#     new_tbl = deepcopy(tbl)
#     p.addnext(new_tbl)



doc.tables[2].cell(0,2).text='报告编号：2342'
paragraph=doc.tables[2].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[4].cell(0,2).text='报告编号：2342'
paragraph=doc.tables[4].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[6].cell(0,2).text='报告编号：2342'
paragraph=doc.tables[6].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.tables[8].cell(0,2).text='报告编号：2342'
paragraph=doc.tables[8].cell(0,2).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.paragraphs[0].text='报告编号：232'
paragraph=doc.paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.tables[3].cell(10,1).text="检验结果见第3页"
paragraph=doc.tables[3].cell(10,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
doc.tables[3].cell(5,3).text='某某检验员'
paragraph=doc.tables[3].cell(5,3).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(4,1).text='--'
paragraph=doc.tables[7].cell(4,1).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.tables[7].cell(4,6).text='--'
paragraph=doc.tables[7].cell(4,6).paragraphs[0]
run=paragraph.runs
font=run[0].font
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# copy_table_after(doc.tables[0],doc.paragraphs[1])



calid1=doc.tables[7].cell(14,5).text
cvn1=doc.tables[7].cell(14,8).text
calid2=doc.tables[7].cell(16,5).text
cvn2=doc.tables[7].cell(16,8).text




def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
table=doc.tables[7]
row1 = doc.tables[7].rows[14]
row2 = doc.tables[7].rows[15]
row3= doc.tables[7].rows[16]
row4 = doc.tables[7].rows[17]

remove_row(table, row1)
remove_row(table, row2)
remove_row(table, row3)
remove_row(table, row4)
count=0
text=[]
for par in doc.paragraphs:
    count+=1
    print(par.text)
    print(count)
    text.append(par.text)
print(doc.paragraphs[60].text)
print(doc.paragraphs[61].text)
print(text)
for item in text:
    if item.startswith('检验日期')==True:
        jianyanshijian=item
    if item.startswith('检验地点')==True:
        jianyandidian=item
print(jianyanshijian)
print(jianyandidian)
def replace_text(doc, old_text, new_text):
    # 遍历每个段落
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if old_text in p.text:
            # 使用 runs 替换内容但不改变样式
            # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
replace_text(doc,'检验人员：','检验人员：李成果')
replace_text(doc,'主检：','主检：李成果')
replace_text(doc,'报告编写人：','报告编写人： 李成果 ')
replace_text(doc,'报告校对人：','报告校对人： 龚香坤 ')


doc.save('修改编号后.docx')


canshuqueren_table=doc.tables[7]


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
def get_paragraph(paras, text):
    for para in paras:
        if text in para.text:
            return para
    raise KeyError("The text cannot be found anywhere in the document")
para1=get_paragraph(doc_canshu.paragraphs,'参数确认表')
move_table_after(canshuqueren_table,para1)


canshuqueren_row1 = doc_canshu.tables[0].rows[0]
canshuqueren_row2 = doc_canshu.tables[0].rows[10]
canshuqueren_row3 = doc_canshu.tables[0].rows[11]
canshuqueren_row4 = doc_canshu.tables[0].rows[12]
canshuqueren_row5 = doc_canshu.tables[0].rows[13]

remove_row(canshuqueren_table, canshuqueren_row1)
remove_row(canshuqueren_table, canshuqueren_row2)
remove_row(canshuqueren_table, canshuqueren_row3)
remove_row(canshuqueren_table, canshuqueren_row4)
remove_row(canshuqueren_table, canshuqueren_row5)

doc_canshu.save('参数确认表111.docx')





