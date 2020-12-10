import docx
def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)
document = docx.Document('bianhao001.docx')
table = document.tables[7]
new_document = docx.Document('参数确认表_模版.docx')
copy_table_after(table,'参数确认')